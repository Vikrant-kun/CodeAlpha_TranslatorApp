import os
import mammoth
from flask import Flask, render_template, request, jsonify
from werkzeug.utils import secure_filename

# --- NEW: Import from deep-translator ---
from deep_translator import GoogleTranslator
from deep_translator.constants import GOOGLE_LANGUAGES_TO_CODES

# --- File Processing & OCR Imports ---
import docx
import PyPDF2
from PIL import Image
import pytesseract

# --- Helper function to get languages in the right format for the frontend ---
def get_languages():
    return {code: name.title() for name, code in GOOGLE_LANGUAGES_TO_CODES.items()}

# --- Upgraded Translation Logic using deep-translator ---
def real_translate(text, target_lang, source_lang='auto'):
    if not text or not text.strip():
        return ""
    try:
        if isinstance(text, list):
            return GoogleTranslator(source=source_lang, target=target_lang).translate_batch(text)
        else:
            return GoogleTranslator(source=source_lang, target=target_lang).translate(text)
    except Exception as e:
        print(f"Translation error with deep-translator: {e}")
        return "Error: Translation failed."

# --- Upgraded DOCX to HTML Translation ---
def translate_docx_to_html(filepath, target_lang):
    try:
        doc = docx.Document(filepath)
        original_texts = [para.text for para in doc.paragraphs if para.text.strip()]
        if not original_texts: return "<p>No text found in document.</p>"
        translated_texts = real_translate(original_texts, target_lang)
        
        translated_iter = iter(translated_texts)
        for para in doc.paragraphs:
            if para.text.strip():
                for i in range(len(para.runs)): para.runs[i].text = ''
                para.runs[0].text = next(translated_iter, '')
                
        result = mammoth.convert_to_html(doc)
        return result.value
    except Exception as e:
        print(f"Error processing docx for preview: {e}")
        return "<h3>Error: Could not generate a preview.</h3>"

# --- File Extraction (for PDF/Images) ---
def extract_text_from_pdf(filepath):
    try:
        text = ""
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                extracted_text = page.extract_text()
                if extracted_text: text += extracted_text
        return text
    except Exception as e: return None

def extract_text_from_image(filepath):
    try: return pytesseract.image_to_string(Image.open(filepath))
    except Exception as e: return None

# --- Flask App ---
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

@app.route("/")
def index():
    return render_template("index.html", languages=get_languages())

@app.route("/translate-text", methods=["POST"])
def translate_text_route():
    data = request.get_json()
    text = data.get('text', '')
    # THE FIX: Now correctly getting source and target from the new frontend
    source = data.get('source_language', 'auto') 
    target = data.get('target_language', 'en')
    translated_text = real_translate(text, target, source)
    return jsonify({'translated_text': translated_text})

@app.route("/translate-file", methods=["POST"])
def translate_file_route():
    if 'file' not in request.files: return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '': return jsonify({'error': 'No file selected'}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    target_language = request.form.get('language', 'en')

    try:
        if filename.lower().endswith('.docx'):
            html_content = translate_docx_to_html(filepath, target_language)
            return jsonify({'file_type': 'docx', 'html_content': html_content})

        text_to_translate = ""
        if filename.lower().endswith('.pdf'):
            text_to_translate = extract_text_from_pdf(filepath)
        elif filename.lower().endswith(('.png', '.jpg', '.jpeg')):
            text_to_translate = extract_text_from_image(filepath)
        else:
            return jsonify({'error': 'Unsupported file type'}), 400
        
        if not text_to_translate: return jsonify({'file_type': 'text', 'translated_text': 'No text could be extracted from this file.'})
        translated_text = real_translate(text_to_translate, target_language)
        return jsonify({'file_type': 'text', 'translated_text': translated_text})
    finally:
        if os.path.exists(filepath): os.remove(filepath)

if __name__ == "__main__":
    app.run(debug=True)

